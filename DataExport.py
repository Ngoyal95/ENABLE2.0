import docx
import time
import pyqtgraph as pg
import pandas as pd
import sys
from PyQt5.QtWidgets import QMessageBox, QDialog
from openpyxl import Workbook
from pprint import pprint

#### EXCEL (SINGLE SPREADSHEETS AND COMPILED) ####
def exportToExcel(StudyRoot,OutDir):
    '''
    Function exports all patient data (for patients who ignore == False) under the StudyRoot to an excel spreadsheet (Compiled data set),
    individual patientexcel spreasheets, and properly compiled values for graph generation
    '''
    colHeaders = ['Exam','Date','Modality','Follow-Up','Name','Tool',\
    'Description','Target','Sub-Type','Series','Slice#','RECIST Diameter (cm)', \
    'Long Diameter (mm)','Short Diameter (mm)','Volume (mm³)','HU Mean (HU)','Creator', \
    'Target RECIST Sum (mm)', 'Best Response Sum (mm)', 'Non-Target RECIST Sum (mm)',\
    'Target RECIST Sum percent change from baseline','Target RECIST Sum percent change from best response',\
    'Non-Target RECIST Sum percent change from baseline','Target Response',\
    'Non-Target Response','Overall Response']

    compileWb = Workbook()
    wsC = compileWb.active #use the active sheet (first and only one in the doc, named 'Sheet')
    compRow = 1 #index used for printing info to compiled doc, does not reset at end of loop
    
    for key, patient in StudyRoot.patients.items():
        #populate single pt sheets and the compiled cohort sheet
        ptWb = Workbook()
        wsP = ptWb.active
        compRow = mapPtDataExcel(wsP,wsC,compRow,patient,colHeaders)
        compRow+=3 #3 row gap btwn patients
        try:
            ptWb.save(OutDir+'/'+patient.name+'.xlsx') #save individual sheet
        except Exception as e:
            QMessageBox.warning(None,'Error!','Could not save Compiled Data spreadsheet because the file is already open. Please close file and repeat spreadsheet export')
            print("Error: ",e)     
    try:
        compileWb.save(OutDir+'/'+'CompiledData.xlsx')
    except PermissionError:
        QMessageBox.warning(None,'Error!','Could not save Compiled Data spreadsheet because the file is already open. Please close file and repeat spreadsheet export')

def mapPtDataExcel(wsP, wsC, compRow, patient, colHeaders):
    '''
    Map data from patient object to a spreadsheet
    '''
    ptRow = 1 #index used for printing single pt sheets
    numRows = len(patient.exams.keys()) #total number of rows for a patient (exams + lesions)
    for key, exam in patient.exams.items():
        if exam.ignore == False:
            numRows += len(exam.lesions) #count rows only for included exams

    ptData = ['Patient name:',patient.name,'MRN:',patient.mrn,'Protocol:',patient.sid]
    #print patient ID info:
    for i in range(0,len(ptData)):
        wsP.cell(row = ptRow, column = i+1).value = ptData[i]
        wsC.cell(row = compRow, column = i+1).value = ptData[i]
    ptRow += 1
    compRow += 1

    #print headers:
    for i in range(0,len(colHeaders)):
        wsP.cell(row = ptRow, column = i+1).value = colHeaders[i]
        wsC.cell(row = compRow, column = i+1).value = colHeaders[i]
    ptRow += 1
    compRow += 1

    #print data:
    for key,exam in patient.exams.items():
        if exam.ignore is False:
            examData = pullExamData(exam, patient)
            col = 3
            #len(examData) #skip to columns where we print lesion data, +1 due to 1 indexing
            examData[0] = key
            for z in range(0,col):
                wsP.cell(row = ptRow, column = z+1).value = examData[z]
                wsC.cell(row = compRow, column = z+1).value = examData[z]

            #print RECIST Data
            for z in range(2,len(examData)):
                wsP.cell(row = ptRow, column = z+15).value = examData[z]
                wsC.cell(row = compRow, column = z+15).value = examData[z]
            ptRow+=1
            compRow+=1

            for lesion in exam.lesions:
                lesionData = pullLesionData(lesion)			
                for z in range(col,col+len(lesionData)):
                    wsP.cell(row = ptRow, column = z+1).value = lesionData[z-col]
                    wsC.cell(row = compRow, column = z+1).value = lesionData[z-col]
                ptRow+=1
                compRow+=1
        elif exam.ignore is True:
            pass
    ptRow = 0
    return compRow

def pullExamData(exam,patient):
    '''
    Extracts exam data and places in list
    '''
    return [exam.exam_num,exam.date,exam.modality,exam.trecistsum,patient.bestresponse,exam.\
            ntrecistsum,exam.tfrombaseline,exam.tfrombestresponse,exam.ntfrombaseline,exam.tresponse,exam.ntresponse,exam.overallresponse]

def pullLesionData(lesion):
    ''''
    Extracts lesion data and place in list
    '''
    return [lesion.fu,lesion.name,lesion.tool,lesion.desc,lesion.target,\
    lesion.subtype,lesion.series,lesion.slice,round(lesion.recistdia/10,1),lesion.longdia,\
    lesion.shortdia,lesion.volume,lesion.humean,lesion.creator]


#### PLOT DATA GEN AND EXPORT ####
def waterfallPlot(StudyRoot):
    #this function will return two lists: one containig plotting data, and a key relating patients to their number

    #Get data
    responses = []
    for key,patient in StudyRoot.patients.items():
        if patient.ignore == False:
            #error handling
            try:
                responses.append(float(((patient.bestresponse-patient.baselinesum)*100)/patient.baselinesum)) #create list of best responses
            except ZeroDivisionError:
                QMessageBox.warning(None,'Error!','Could not compute percent change for patient: ' + patient.name + '\nError due to baseline sum = 0. Patient will NOT be plotted and NOT exported.')
    
    responses.sort(reverse = True)
    return responses

def spiderPlot(StudyRoot):
    '''
    Return data for spider plot, dict of patients with values as tuple w/ format (lists of their target RECIST sum,weeks from baseline,days from baseline)
    '''
    ptData = {}
    popPat = False
    for key,patient in StudyRoot.patients.items():
        popPat = False
        if patient.ignore == False:
            meas = []
            deltaTW = []
            deltaTD = []

            for key2,exam in patient.exams.items():
                try:
                    meas.append(exam.trecistsum/patient.baselinesum)
                    deltaTW.append(exam.weeksfromB)
                    deltaTD.append(exam.daysfromB)				
                except ZeroDivisionError:
                    popPat = True #dont include the patient because of bad baseline
                    meas.append(0)
                    deltaTW.append(exam.weeksfromB)
                    deltaTD.append(exam.daysfromB)
                    #QMessageBox.warning(None,'Error!','Could not compute percent change for patient: ' + patient.name + '\nError due to baseline sum = 0')
            if popPat == False:
                ptData[key] = [meas,deltaTW,deltaTD]
    
    return ptData

def swimmerPlot(StudyRoot):
    #TODO
    pass

def exportPlotData(StudyRoot,OutDir):
    '''
    Exports data for all 3 types of plots (waterfall, spider, swimmer) to an Excel spreadsheet
    '''
    dataWb = Workbook()
    WF = dataWb.active
    WF.title = "Waterfall_Data" #change name first sheet in workbook

    ##### ---- Waterfall ---- #####
    #Column headers
    WFHeaders = ['Patient Number','Patient','MRN','Protocol',r'Best response % change from Baseline','Overall RECIST Response','Clinical Response']
    row = 1
    for i in range(0,len(WFHeaders)):
        WF.cell(row = row, column = i+1).value = WFHeaders[i]
    row+=1

    ptData = []
    for key, patient in StudyRoot.patients.items():
        try:
            val = float(((patient.bestresponse-patient.baselinesum)*100)/patient.baselinesum)
        except ZeroDivisionError:
            val = None
            QMessageBox.warning(None,'Error!','Could not compute percent change for patient: ' + patient.name + '\nError due to baseline sum = 0')

        indivPtData = [patient.name,int(patient.mrn),patient.sid,val,patient.exams[1].overallresponse]
        ptData.append(indivPtData)

    #now reorder patients in ptData list descending %change from baseline. Items with None as percent change go to bottom
    ptData.sort(key = lambda x: (x[3] is None, x[3]),reverse=False) #increasing order, NoneType at end
    temp = [x for x in ptData if x[3] is not None]
    temp.sort(key = lambda x: x[3],reverse=True)
    ptData = temp + [x for x in ptData if x[3] is None] #Split and reconcat to get proper ordering

    #print data and enumerate patients
    for i in range(0,len(ptData)):
        for k in range(0,len(ptData[i])):
                WF.cell(row = row, column = k+2).value = ptData[i][k] #note, +2 used to avoid the Patient Number col
                WF.cell(row = row, column = 1).value = i+1 #enumerate patients, index at 1
        row+=1

    #### ---- Spider ---- ####
    #### ---- Swimmer ---- ####

    try:
        dataWb.save(OutDir+'/'+'plottingData.xlsx')
    except PermissionError:
        QMessageBox.warning(None,'Error!','Could not save Plot Data spreadsheet because the file is already open. Please close file and spreadsheet export')
    
#### LABMATRIX ####
def export_to_labmatrix(StudyRoot,OutDir,LMClientPath,user,pw):
    '''
    Function to format data for labmatrix and then upload using RadioglogyImportClient.jar.
    '''
    LMHeaders = []
    time = time.strftime("%H:%M")
    day = time.strftime("%d/%m/%Y")
    LMF = OutDir + "/" + day + "_" + time #Folder to save LM spreadsheets to
    if not os.path.exists(LMF):
        os.makedirs(LMF)

    for key,patient in StudyRoot.items():
        ptWb = Workbook()
        wsP = ptWb.active
        map_labmatrix_data_models(wsP,patient,LMHeaders)
        savePath = LMF + '/' + patient.name + '.xlsx'
        ptWb.save(LMF + '/' + patient.name + '.xlsx') #save
        os.system("java -cp" + LMClientPath + "ImportArgs" + " " + user + " " + pw + " " + savePath) #upload

def map_labmatrix_data_models(wsP,patient,LMHeaders):
    '''
    Produce data model for upload to Labmatrix
    '''
    ptRow = 1 #row index for current document
    for i in range(0,len(ptData)):
        wsP.cell(row = ptRow, column = i+1).value = ptData[i]
        wsC.cell(row = compRow, column = i+1).value = ptData[i]
    ptRow += 1
    compRow += 1

#### CONSULT LOG ####
def exportToLog(RECISTDir,OutDir,StudyRoot,vals):
    '''
    Create radiologist consultation log.
    Expected format of vals is list of the following: [selkey,consultant,phys,date,priorbaseline,baseline,restaging,describe_1,reason_2a,reason_2b,describe_2,comments,time]
    '''
    patient = StudyRoot.patients[vals[0]]
    template = docx.Document(RECISTDir+r'/CIPS_Consult_Log.docx')
    
    table = template.tables[0] #meeting log table
    table.cell(0,3).text = patient.sid
    table.cell(1,1).text = vals[2]
    table.cell(1,3).text = vals[3]
    table.cell(2,1).text = patient.name
    table.cell(2,3).text = patient.mrn

    table = template.tables[1] #dates table
    table.cell(0,1).text = vals[3]
    table.cell(0,3).text = vals[4]
    table.cell(1,3).text = vals[5]
    table.cell(2,3).text = vals[6]

    table = template.tables[2] #reason for review
    if vals[7] is not '':
        table.cell(0,0).text = 'X'
        table.cell(1,3).text = vals[7]
    if vals[8] or vals[9]:
            table.cell(3,0).text = 'X'
            if vals[8]:
                table.cell(5,1).text = 'X'
                table.cell(6,3).text = vals[10]
            elif vals[9]:
                table.cell(7,1).text = 'X'
                table.cell(8,3).text = vals[10]
    
    #Use restaging date to extract exam data
    exam = next((x for key,x in patient.exams.items() if x.date == vals[6]), None)
    
    table = template.tables[3]
    table.cell(0,1).text = exam.modality + ' - ' + exam.description
    table.cell(0,4).text = exam.date

    table = template.tables[4] #lesion data table
    #count number T,NT,NL lesions
    lesionCount = 0
    for lesion in exam.lesions:
        if lesion.params['Target'].lower() != 'unspecified':
            lesionCount += 1
    
    for i in range(0,lesionCount):
        lesionData = [
                    exam.lesions[i].params['Target'], exam.modality, exam.lesions[i].params['Description'], \
                    round(exam.lesions[i].params['RECIST Diameter (mm)']/10,2), exam.lesions[i].params['Series'], exam.lesions[i].params['Slice#']]
        for l in range(0,6):
            col = l+1
            row = i+1
            table.cell(row,col).text = str(lesionData[l])
    
    template.tables[5].cell(0,0).text = vals[11] #comments
    
    #save, throw error if open
    try:
        template.save(OutDir + r'/TEMPLATE.docx')
    except Exception as e:
        print('Error:', e)
    


    