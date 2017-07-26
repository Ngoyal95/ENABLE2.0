#RECIST Worksheet generator module
#Revision 6/21/17
import docx
import re
import os
import pathlib

def generate_recist_sheets(RECISTDir, OutDir, StudyRoot,singleSheet):
    '''
    Generate the RECIST sheets for each patient in the study (if they have not been put in the excludedlist)
    '''
    for key,patient in StudyRoot.patients.items():
        if patient.ignore == False:
            if singleSheet == True:
                numExams = len(patient.exams.keys())
                for l in range(1, numExams+1):
                    #search for current exam
                    if patient.exams[l].current == True:
                        break
                recist_sheet(RECISTDir, OutDir, patient, patient.exams[l]) #generate sheet for exam marked as 'current'
            elif singleSheet == False:
                #iterate through all exams within the baseline->current range and create RECIST worksheets
                for key, exam in patient.exams.items():
                    if exam.ignore == False:
                        recist_sheet(RECISTDir, OutDir, patient, exam)
    
def get_mrn_sid(file):
    '''
    Extract the MRN and SID from the BL filename (expected format is MRN#######_##-X-####)
    '''
    regMRN = re.compile(r'\d{7}')
    regSID = re.compile(r'\w{2}-\w-\w{4}')
    MRN = regMRN.search(file).group()
    SID = regSID.search(file).group()
    return MRN,SID

def convert_date(date):
    '''
    converts date from format mm/dd/yyy to mm.dd.yyyy
    '''
    return str(date.replace(r'/','.',))

def recist_sheet(RECISTDir, OutDir, patient, exam):
    '''
    Generate a single RECIST sheet for the patient and specified exam
    '''
    MRN = patient.mrn
    SID = patient.study_protocol
    
    #### GENERAL DATA ####
    template = docx.Document(RECISTDir + '\\RECISTForm.docx')
    table = template.tables[4]
    table.cell(1,0).text = patient.name + '\n' + MRN
    table = template.tables[1] #print study info
    table.cell(0,1).text = SID
    if (exam.baseline == True): #based on current exam
        table.cell(0,5).text = 'X'
    else:
        table.cell(0,8).text = 'X'
    table.cell(0,3).text = 'Course ' + patient.course + '  /Day ' + patient.day

    #### TUMOR MEASUREMENT DATA ####
    table = template.tables[2]
    totNumLesion = len(exam.lesions)
    numLesion = 0
    for lesion in exam.lesions:
        if lesion.params['Target'].lower() != 'unspecified':
            numLesion += 1
    for i in range(0,numLesion):
        if i > 11:
            break # only 12 lesion entry rows in RECIST sheet
        row = i+1
        table.cell(row,1).text = str(exam.lesions[i].params['Description'])
        table.cell(row,2).text = str(exam.lesions[i].params['Target'])
        if exam.lesions[i].newlesion == True:
            table.cell(row,3).text = "New Lesion"
        table.cell(row,4).text = exam.modality
        table.cell(row,6).text = (str(exam.lesions[i].params['Series']) + '/' + str(exam.lesions[i].params['Slice#']))
        table.cell(row,7).text = str(round(exam.lesions[i].params['RECIST Diameter (mm)']/10,1)) #convert to cm, round to tenths
        table.cell(row,8).text = exam.date

    #### RESPONSE DATA ####
    table = template.tables[3]
    table.cell(0,1).text = str(exam.trecistsum)
    table.cell(1,1).text = str(patient.baselinesum)
    table.cell(2,1).text = str(patient.bestresponse)
    table.cell(3,1).text = str(exam.tfrombestresponse)
    table.cell(4,1).text = str(exam.tfrombaseline)
    table.cell(5,1).text = exam.tresponse
    #table.cell(6,1).text = exam.ntresponse
    #table.cell(7,1).text = exam.overallresponse

    #### Date measured, by who ####
    table = template.tables[5]
    table.cell(0,1).text = exam.date
    table.cell(1,1).text = exam.measuredby

    #********** Temporary fix for the font size of the output RECIST worksheets ******** MUST CHANGE IN FUTURE#
    for table in template.tables[1:5]:
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= docx.shared.Pt(7)
                        
    path = pathlib.Path(OutDir + r'/MRN' + MRN + "_" + SID ).mkdir(parents=True, exist_ok=True) #check if path exists (individual pt. folder) if not, create
    try:
        template.save(OutDir + r'/MRN' + MRN + "_" + SID +  r'/MRN' + MRN + "_" + SID + "_Exam_" + convert_date(exam.date) + ".docx") #save word document
    except Exception as e:
        #file is open
        print('Error: ', e)